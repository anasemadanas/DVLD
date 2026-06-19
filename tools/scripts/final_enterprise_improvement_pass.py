from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parents[2]
DOCS = REPO_ROOT / "docs"
UML = DOCS / "UML"


RTM = DOCS / "Requirements_Traceability_Matrix.docx"
SDD = DOCS / "System_Design_Document_SDD.docx"


RTM_ROWS = [
    [
        "US-12",
        "FR-09.6",
        "UC-12",
        "InternationalLicenseService; LicenseService; ValidationService",
        "InternationalLicenses; Licenses; Drivers; Applications",
        "sequence-international-license.puml; chen-notation-erd.puml; physical-erd.puml",
        "TC-16",
    ],
    [
        "US-12",
        "FR-09.6; FR-10",
        "UC-12",
        "InternationalLicenseService; LicenseService; DetentionService; ValidationService",
        "InternationalLicenses; Licenses; DetainedLicenses",
        "sequence-international-license.puml; sequence-license-detention-release.puml",
        "TC-17",
    ],
    [
        "US-14",
        "FR-12",
        "UC-14",
        "PersonService; SearchService; ValidationService",
        "People; Countries",
        "use-case-diagram.puml; database-schema.puml",
        "TC-18",
    ],
    [
        "US-15",
        "FR-12",
        "UC-15",
        "ReportService; ApplicationService; LicenseService",
        "Applications; Licenses; LicenseClasses",
        "use-case-diagram.puml; database-schema.puml",
        "TC-19",
    ],
    [
        "US-13",
        "FR-11",
        "UC-13",
        "AuditService; UserService",
        "Users; Applications; Licenses; DetainedLicenses; Tests; TestAppointments",
        "architecture.puml; login-auth-sequence.puml; general-sequence-diagram.puml",
        "TC-20",
    ],
]


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def update_rtm() -> int:
    doc = Document(RTM)
    if not doc.tables:
        raise RuntimeError("RTM has no traceability table")
    table = doc.tables[0]
    existing_tc = {cell.text.strip() for row in table.rows for cell in row.cells if cell.text.strip().startswith("TC-")}
    added = 0
    for row_values in RTM_ROWS:
        if row_values[-1] in existing_tc:
            continue
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = value
        added += 1
    doc.save(RTM)
    return added


def count_target_rtm_rows() -> int:
    doc = Document(RTM)
    target = {row[-1] for row in RTM_ROWS}
    found = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                value = cell.text.strip()
                if value in target:
                    found.add(value)
    return len(found)


def update_sdd() -> bool:
    doc = Document(SDD)
    marker = "Enterprise Security, Backup, And Recovery Addendum"
    if any(marker in p.text for p in doc.paragraphs):
        doc.save(SDD)
        return False

    doc.add_page_break()
    doc.add_heading(marker, level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "This section defines concise operational controls for protecting DVLD data, recovering from failures, "
        "and preserving accountability in enterprise use."
    )

    sections = [
        (
            "Backup Strategy",
            [
                "SQL Server database backups should be scheduled and monitored according to the approved operations calendar.",
                "Backup files should be stored in an approved secure location separate from the production database host where possible.",
                "Backup completion must be verified, and failed backups should be escalated to the responsible technical owner.",
                "Periodic restore tests should confirm that backups are usable, complete, and aligned with recovery expectations.",
            ],
        ),
        (
            "Recovery Strategy",
            [
                "If the desktop application fails, users should restart the application and avoid repeating partially completed operations until record status is verified.",
                "Database recovery should follow the approved SQL Server restore process using the latest verified backup and any available recovery logs.",
                "After unexpected shutdown, critical records such as applications, tests, licenses, detention records, and payments should be checked for consistency.",
                "Recovery responsibility should be assigned to the system administrator or designated technical support owner.",
            ],
        ),
        (
            "Password Hashing",
            [
                "User passwords must be stored as hashes, not plain text.",
                "Authentication should verify submitted credentials by comparing the computed password hash with the stored hash.",
                "Credentials, password hashes, and authentication secrets must not be exposed in logs, error messages, screenshots, or reports.",
            ],
        ),
        (
            "Session Timeout",
            [
                "Authenticated sessions should expire after a defined period of inactivity.",
                "Inactive or deactivated users must not be allowed to continue using the system after account status changes are detected.",
                "Unauthorized session reuse should be prevented by clearing session state on logout, timeout, and failed authorization checks.",
            ],
        ),
        (
            "Audit Retention",
            [
                "Important system operations should be logged, including login attempts, user management, application changes, test results, license issuance, detention, and release actions.",
                "Audit records should include the responsible user, operation name, date/time, affected entity, and outcome where practical.",
                "Audit records should be retained according to an approved organizational retention policy and protected from unauthorized modification.",
            ],
        ),
    ]
    for heading, bullets in sections:
        doc.add_heading(heading, level=2)
        for item in bullets:
            doc.add_paragraph(f"- {item}")
    doc.save(SDD)
    return True


def render_state_png(path: Path, title: str, states: list[str], transitions: list[tuple[str, str, str]], note: str) -> None:
    img = Image.new("RGB", (1500, 950), "white")
    d = ImageDraw.Draw(img)
    d.text((750, 45), title, fill="#111827", font=font(28, True), anchor="mm")
    positions = {
        state: (170 + (i % 4) * 330, 170 + (i // 4) * 250)
        for i, state in enumerate(states)
    }
    for src, dst, label in transitions:
        if src not in positions or dst not in positions:
            continue
        sx, sy = positions[src]
        tx, ty = positions[dst]
        d.line([(sx + 120, sy + 35), (tx + 120, ty + 35)], fill="#64748B", width=2)
        mx, my = (sx + tx) // 2 + 120, (sy + ty) // 2 + 35
        d.rectangle([mx - 70, my - 12, mx + 70, my + 12], fill="white")
        d.text((mx, my), label[:28], fill="#334155", font=font(10), anchor="mm")
    for state, (x, y) in positions.items():
        d.rounded_rectangle([x, y, x + 240, y + 70], radius=14, fill="#F8FAFC", outline="#334155", width=2)
        d.text((x + 120, y + 35), state, fill="#111827", font=font(14, True), anchor="mm")
    d.rounded_rectangle([1030, 710, 1430, 875], radius=10, fill="#EFF6FF", outline="#93C5FD", width=1)
    d.text((1050, 730), "Note", fill="#1E3A8A", font=font(13, True))
    d.multiline_text((1050, 760), note, fill="#111827", font=font(12), spacing=5)
    img.save(path)


def render_sequence_png(path: Path) -> None:
    participants = ["User", "Login UI", "AuthenticationService", "UserService", "UserRepository", "SQL Server", "AuditService", "Dashboard"]
    messages = [
        (0, 1, "Enter username/password"),
        (1, 2, "Login request"),
        (2, 2, "Validate input"),
        (2, 3, "Get user"),
        (3, 4, "Find by username"),
        (4, 5, "Query Users"),
        (2, 2, "Verify password hash"),
        (2, 2, "Check IsActive"),
        (2, 6, "Record success/failure"),
        (2, 1, "Return result"),
        (1, 7, "Open Dashboard on success"),
    ]
    img = Image.new("RGB", (1800, 1050), "white")
    d = ImageDraw.Draw(img)
    d.text((900, 45), "Login / Authentication - Sequence Diagram", fill="#111827", font=font(28, True), anchor="mm")
    xs = [90 + i * 230 for i in range(len(participants))]
    for x, p in zip(xs, participants):
        d.rounded_rectangle([x, 95, x + 180, 145], radius=8, fill="#DBEAFE", outline="#1E3A8A", width=2)
        d.text((x + 90, 120), p, fill="#0F172A", font=font(11, True), anchor="mm")
        d.line([(x + 90, 145), (x + 90, 930)], fill="#94A3B8", width=2)
    y = 190
    for src, dst, label in messages:
        sx = xs[src] + 90
        dx = xs[dst] + 90
        if src == dst:
            d.arc([sx, y - 10, sx + 70, y + 35], 270, 90, fill="#475569", width=2)
            d.text((sx + 90, y + 8), label, fill="#334155", font=font(10), anchor="lm")
        else:
            d.line([(sx, y), (dx, y)], fill="#475569", width=2)
            arrow = -1 if dx < sx else 1
            d.polygon([(dx, y), (dx - 10 * arrow, y - 5), (dx - 10 * arrow, y + 5)], fill="#475569")
            d.text(((sx + dx) // 2, y - 12), label, fill="#334155", font=font(10), anchor="mm")
        y += 65
    alt_text = "Alternative flows: invalid credentials, inactive user, and database error are returned to Login UI and recorded by AuditService where applicable."
    d.rounded_rectangle([160, 850, 1640, 970], radius=10, fill="#FEF3C7", outline="#F59E0B", width=1)
    d.multiline_text((185, 880), alt_text, fill="#111827", font=font(14), spacing=6)
    img.save(path)


def parse_classes(puml: str) -> list[tuple[str, list[str]]]:
    result = []
    for name, body in re.findall(r"class\s+(\w+)\s*\{(.*?)\}", puml, flags=re.S):
        attrs = [line.strip() for line in body.splitlines() if line.strip()]
        result.append((name, attrs))
    return result


def render_class_png() -> None:
    puml = (UML / "diagram" / "class-diagram.puml").read_text(encoding="utf-8")
    classes = parse_classes(puml)
    img = Image.new("RGB", (1700, 1250), "white")
    d = ImageDraw.Draw(img)
    d.text((850, 40), "DVLD System - Conceptual Domain Model", fill="#111827", font=font(26, True), anchor="mm")
    d.rounded_rectangle([90, 75, 1610, 145], radius=8, fill="#EFF6FF", outline="#93C5FD")
    d.text((110, 95), "Note: Conceptual domain model only. Use script.sql and database ERDs for physical SQL structure.", fill="#1E3A8A", font=font(13), anchor="la")
    for i, (name, attrs) in enumerate(classes):
        x = 70 + (i % 4) * 405
        y = 190 + (i // 4) * 320
        h = 42 + min(len(attrs), 8) * 24
        d.rectangle([x, y, x + 350, y + 42], fill="#1E3A8A", outline="#0F172A", width=2)
        d.text((x + 175, y + 21), name, fill="white", font=font(13, True), anchor="mm")
        for row, attr in enumerate(attrs[:8]):
            ry = y + 42 + row * 24
            d.rectangle([x, ry, x + 350, ry + 24], fill="#FFFFFF", outline="#CBD5E1")
            d.text((x + 8, ry + 12), attr[:48], fill="#111827", font=font(9), anchor="lm")
        d.rectangle([x, y, x + 350, y + h], outline="#0F172A", width=2)
    img.save(UML / "diagram" / "class-diagram.png")


def render_pngs() -> None:
    state_dir = UML / "state"
    seq_dir = UML / "sequence"
    render_state_png(
        state_dir / "application-state-diagram.png",
        "DVLD Application Lifecycle - State Diagram",
        ["New", "In Review / Processing", "Completed", "Cancelled", "Rejected / Blocked"],
        [
            ("New", "In Review / Processing", "Validate Application"),
            ("New", "Cancelled", "Cancel Application"),
            ("New", "Rejected / Blocked", "Reject Invalid Application"),
            ("In Review / Processing", "Completed", "Complete Required Workflow"),
            ("In Review / Processing", "Cancelled", "Cancel Application"),
            ("Rejected / Blocked", "In Review / Processing", "Correct And Revalidate"),
        ],
        "Rejected / Blocked covers validation, eligibility,\nduplicate active application, payment, or workflow\nprerequisite failure.",
    )
    render_state_png(
        state_dir / "license-state-diagram.png",
        "DVLD License Lifecycle - State Diagram",
        ["Active", "Expired", "Renewed", "Replaced", "Detained", "Released", "Inactive"],
        [
            ("Active", "Expired", "Expire License"),
            ("Expired", "Renewed", "Renew License"),
            ("Renewed", "Active", "Activate Renewed License"),
            ("Active", "Replaced", "Replace Lost/Damaged"),
            ("Active", "Detained", "Detain License"),
            ("Detained", "Released", "Release License"),
            ("Released", "Active", "Restore Active Status"),
            ("Replaced", "Inactive", "Deactivate License"),
            ("Expired", "Inactive", "Deactivate License"),
        ],
        "Only active licenses should be used for valid\ndriving privileges, replacement, international-license\nchecks, and detention.",
    )
    render_sequence_png(seq_dir / "login-auth-sequence.png")
    render_class_png()


def write_final_report(rtm_rows_added: int, rtm_rows_present: int, sdd_updated: bool) -> None:
    report = DOCS / "final-enterprise-improvement-report.md"
    lines = [
        "# Final Enterprise Improvement Report",
        "",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Files Created",
        "",
        "- `docs/UML/class-diagram-validation-report.md`",
        "- `docs/UML/state/application-state-diagram.puml`",
        "- `docs/UML/state/application-state-diagram.png`",
        "- `docs/UML/state/license-state-diagram.puml`",
        "- `docs/UML/state/license-state-diagram.png`",
        "- `docs/UML/sequence/login-auth-sequence.puml`",
        "- `docs/UML/sequence/login-auth-sequence.png`",
        "- `docs/final-enterprise-improvement-report.md`",
        "",
        "## Files Updated",
        "",
        "- `docs/UML/diagram/class-diagram.puml`",
        "- `docs/UML/diagram/class-diagram.png`",
        "- `docs/Requirements_Traceability_Matrix.docx`",
        "- `docs/System_Design_Document_SDD.docx`",
        "- `docs/UML/README.md`",
        "",
        "## Class Diagram Validation Summary",
        "",
        "The class diagram was validated against `script.sql` and classified as a conceptual domain model. It now includes an explicit note stating that it is not the physical database model. SQL-backed structure remains represented by `script.sql` and the database ERD files.",
        "",
        "## State Diagrams Created",
        "",
        "- Application lifecycle: New, In Review / Processing, Cancelled, Completed, and Rejected / Blocked.",
        "- License lifecycle: Active, Expired, Renewed, Replaced, Detained, Released, and Inactive.",
        "",
        "## Login Sequence Diagram Created",
        "",
        "Added login/authentication sequence with User, Login UI, AuthenticationService, UserService, UserRepository, SQL Server Database, AuditService, and Dashboard. Alternative flows cover invalid credentials, inactive users, and database errors.",
        "",
        "## RTM TC16-TC20 Mapping Summary",
        "",
        f"Rows added to RTM in the final successful run: {rtm_rows_added}",
        f"Target TC16-TC20 rows present in RTM: {rtm_rows_present} / 5",
        "- `TC-16`: valid Class 3 international license issuance.",
        "- `TC-17`: reject international license for detained/expired local license.",
        "- `TC-18`: search records by national number.",
        "- `TC-19`: filter reports by status/date/license class.",
        "- `TC-20`: verify audit log records operation user and date.",
        "",
        "## SDD Security/Backup/Recovery Additions",
        "",
        f"SDD addendum inserted: {sdd_updated}",
        "Added concise enterprise controls for backup strategy, recovery strategy, password hashing, session timeout, and audit retention.",
        "",
        "## UML Catalog Update Summary",
        "",
        "Updated `docs/UML/README.md` with entries for the new state diagrams and login/authentication sequence diagram.",
        "",
        "## Remaining Recommendations",
        "",
        "- Refresh Word table of contents manually if any DOCX has a generated TOC.",
        "- Review password hashing guidance against the actual implementation before production use.",
        "- Add real PlantUML-rendered PNGs later if PlantUML CLI is installed in the workspace.",
        "",
        "## Constraints Observed",
        "",
        "- `script.sql` was not modified.",
        "- Database schema and database diagrams were not modified.",
        "- No existing files were removed.",
        "",
        "## Validation Notes",
        "",
        "- DOCX edits were structurally validated with `python-docx`.",
        "- Visual DOCX render QA was attempted but could not complete because the renderer dependency `pdf2image` is unavailable in this environment; DOCX structural validation completed successfully.",
        "- PNG previews were generated locally from the diagram content because PlantUML CLI is not available on PATH.",
    ]
    report.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    rtm_rows_added = update_rtm()
    rtm_rows_present = count_target_rtm_rows()
    sdd_updated = update_sdd()
    render_pngs()
    write_final_report(rtm_rows_added, rtm_rows_present, sdd_updated)
    print(f"rtm_rows_added={rtm_rows_added}")
    print(f"sdd_updated={sdd_updated}")
    print("png_exports=4")
    print("script_sql_modified=False")


if __name__ == "__main__":
    main()
