from __future__ import annotations

import re
import shutil
import tempfile
from datetime import datetime
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[2]
DOCS = REPO_ROOT / "docs"
FRS_PATH = DOCS / "Functional_Requirements_Specification_FRS.docx"
RTM_PATH = DOCS / "Requirements_Traceability_Matrix.docx"
SDD_PATH = DOCS / "System_Design_Document_SDD.docx"
REPORT_PATH = DOCS / "final-documentation-quality-check-report.md"


RTM_HEADERS = [
    "User Story ID",
    "Functional Requirement ID",
    "Use Case ID",
    "Related Service",
    "Related Database Table(s)",
    "Related UML Diagram(s)",
    "Related Test Case ID",
]

RTM_ROWS = [
    [
        "US-01",
        "FR-01",
        "UC-01",
        "AuthenticationService; UserService",
        "Users; People",
        "login-auth-sequence.puml; use-case-diagram.puml",
        "TC-01",
    ],
    [
        "US-01",
        "FR-01",
        "UC-01",
        "AuthenticationService; UserService",
        "Users",
        "login-auth-sequence.puml; general-sequence-diagram.puml",
        "TC-02",
    ],
    [
        "US-02",
        "FR-02",
        "UC-02",
        "PersonService; CountryService",
        "People; Countries",
        "use-case-diagram.puml; database-schema.puml",
        "TC-03",
    ],
    [
        "US-02",
        "FR-02",
        "UC-02",
        "PersonService; CountryService",
        "People; Countries",
        "use-case-diagram.puml; database-schema.puml",
        "TC-04",
    ],
    [
        "US-03",
        "FR-04; FR-05",
        "UC-03",
        "ApplicationService; ApplicationTypeService; LicenseClassService",
        "Applications; LocalDrivingLicenseApplications; ApplicationTypes; LicenseClasses",
        "activity-license-issuance.puml; database-overview.puml",
        "TC-05",
    ],
    [
        "US-03",
        "FR-04; FR-05",
        "UC-03",
        "ApplicationService; ApplicationTypeService; LicenseClassService",
        "Applications; LocalDrivingLicenseApplications; ApplicationTypes; LicenseClasses",
        "activity-license-issuance.puml; database-schema.puml",
        "TC-06",
    ],
    [
        "US-04",
        "FR-07",
        "UC-04",
        "TestService; TestTypeService; ApplicationService",
        "TestAppointments; TestTypes; LocalDrivingLicenseApplications; Applications",
        "sequence-test-management.puml; activity-retake-test.puml",
        "TC-07",
    ],
    [
        "US-05",
        "FR-07",
        "UC-05",
        "TestService; AuditService",
        "Tests; TestAppointments; TestTypes; Users",
        "sequence-test-management.puml; general-sequence-diagram.puml",
        "TC-08",
    ],
    [
        "US-06",
        "FR-07; FR-09.1",
        "UC-06",
        "LicenseService; TestService; ApplicationService",
        "Licenses; Tests; TestAppointments; Applications",
        "issue-driving-license-sequence.puml; license-issuance-activity.puml",
        "TC-09",
    ],
    [
        "US-06",
        "FR-03; FR-05; FR-09.1",
        "UC-06",
        "LicenseService; DriverService; TestService; AuditService",
        "Drivers; Licenses; Applications; Tests; LicenseClasses; Users",
        "issue-driving-license-sequence.puml; license-issuance-activity.puml",
        "TC-10",
    ],
    [
        "US-07",
        "FR-09.2",
        "UC-07",
        "LicenseService; TestService",
        "Licenses; Applications; Tests",
        "sequence-license-renewal.puml; activity-license-renewal.puml",
        "TC-11",
    ],
    [
        "US-08",
        "FR-09.3",
        "UC-08",
        "LicenseService; AuditService",
        "Licenses; Applications; DetainedLicenses; Users",
        "sequence-lost-license-replacement.puml; activity-lost-license-replacement.puml",
        "TC-12",
    ],
    [
        "US-09",
        "FR-09.4",
        "UC-09",
        "LicenseService; AuditService",
        "Licenses; Applications; Users",
        "sequence-damaged-license-replacement.puml; activity-damaged-license-replacement.puml",
        "TC-13",
    ],
    [
        "US-10",
        "FR-10",
        "UC-10",
        "DetentionService; LicenseService; AuditService",
        "DetainedLicenses; Licenses; Users",
        "sequence-license-detention.puml; activity-license-detention-release.puml",
        "TC-14",
    ],
    [
        "US-11",
        "FR-09.5; FR-10",
        "UC-11",
        "DetentionService; LicenseService; ApplicationService; AuditService",
        "DetainedLicenses; Licenses; Applications; Users",
        "sequence-license-release.puml; activity-license-detention-release.puml",
        "TC-15",
    ],
    [
        "US-12",
        "FR-09.6",
        "UC-12",
        "InternationalLicenseService; LicenseService; DriverService; AuditService",
        "InternationalLicenses; Licenses; Drivers; Applications; Users",
        "sequence-international-license.puml; activity-international-license.puml",
        "TC-16",
    ],
    [
        "US-12",
        "FR-09.6",
        "UC-12",
        "InternationalLicenseService; LicenseService; DetentionService",
        "InternationalLicenses; Licenses; Drivers; DetainedLicenses",
        "sequence-international-license.puml; sequence-license-detention-release.puml",
        "TC-17",
    ],
    [
        "US-13",
        "FR-02; FR-09.7",
        "UC-13",
        "PersonService; ApplicationService; DriverService; LicenseService",
        "People; Countries; Applications; Drivers; Licenses",
        "use-case-diagram.puml; database-schema.puml",
        "TC-18",
    ],
    [
        "US-15",
        "FR-12",
        "UC-13",
        "ApplicationService; LicenseService; LicenseClassService",
        "Applications; Licenses; LicenseClasses",
        "system-inputs.puml; ui-flow-navigation.puml",
        "TC-19",
    ],
    [
        "US-14",
        "FR-11",
        "UC-03; UC-06; UC-10; UC-11",
        "AuditService; UserService",
        "Users; Applications; Licenses; DetainedLicenses; Tests; TestAppointments",
        "architecture.puml; login-auth-sequence.puml; general-sequence-diagram.puml",
        "TC-20",
    ],
]


def replace_cell_text(cell, text: str) -> None:
    cell.text = text


def normalize_requirement_sentence(text: str) -> str:
    text = text.strip()
    text = text.replace("view User information", "view user information")
    text = text.replace("The system shall provide the ability to manage system users", "The system shall allow administrators to manage system users")
    text = text.replace("The system shall provide the ability to manage persons", "The system shall allow users to manage persons")
    text = text.replace("The system shall provide the ability to manage applications", "The system shall allow users to manage applications")
    if text and not text.endswith("."):
        text += "."
    return text


def update_frs() -> list[str]:
    doc = Document(FRS_PATH)
    fixes = []
    for table in doc.tables:
        if len(table.columns) != 2:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header != ["ID", "Requirement"]:
            continue
        for row in table.rows[1:]:
            original = row.cells[1].text
            updated = normalize_requirement_sentence(original)
            if updated != original:
                replace_cell_text(row.cells[1], updated)
                fixes.append(f"FRS: normalized wording/punctuation for {row.cells[0].text.strip()}.")
    doc.save(FRS_PATH)
    return fixes


def update_rtm() -> list[str]:
    doc = Document(RTM_PATH)
    table = doc.tables[0]
    while len(table.rows) < len(RTM_ROWS) + 1:
        table.add_row()
    while len(table.rows) > len(RTM_ROWS) + 1:
        table._tbl.remove(table.rows[-1]._tr)

    for idx, header in enumerate(RTM_HEADERS):
        replace_cell_text(table.rows[0].cells[idx], header)
    for row_idx, row_data in enumerate(RTM_ROWS, start=1):
        for col_idx, value in enumerate(row_data):
            replace_cell_text(table.rows[row_idx].cells[col_idx], value)
    doc.save(RTM_PATH)
    return [
        "RTM: rebuilt the 20 traceability rows so TC-01 through TC-20 map to the matching QA test scenarios.",
        "RTM: replaced service names not present in the SDD service layer with existing documented services.",
        "RTM: corrected TC-16 through TC-20 user story, FR, use case, service, table, and UML mappings.",
    ]


def update_sdd_tables() -> list[str]:
    doc = Document(SDD_PATH)
    fixes = []

    security_table = None
    for table in doc.tables:
        if len(table.rows) >= 2 and table.rows[0].cells[0].text.strip() == "Area" and table.rows[0].cells[1].text.strip() == "Security Consideration":
            security_table = table
            break
    if security_table:
        seen_session = 0
        for row in security_table.rows[1:]:
            if row.cells[0].text.strip() == "Session Management":
                seen_session += 1
                if seen_session == 2:
                    replace_cell_text(row.cells[0], "Session Protection")
                    fixes.append("SDD: renamed duplicate security row from Session Management to Session Protection.")

    doc.save(SDD_PATH)
    return fixes


def patch_sdd_textbox() -> list[str]:
    old = "The Business Layer contains services responsible for business rules and operational workflows, including application types, license classes, test types, countries, and audit logging services."
    new = "The Presentation Layer is responsible for rendering the graphical user interface and handling user interactions. It forwards user requests to the Business Layer and displays processed results."
    heading_old = '<w:t xml:space="preserve">3.3 Class Diagram </w:t>'
    heading_new = "<w:t>3.3 Conceptual Class Diagram</w:t>"
    with ZipFile(SDD_PATH, "r") as zin:
        document_xml = zin.read("word/document.xml").decode("utf-8")
        fixes = []
        old_count = document_xml.count(old)
        if old_count:
            document_xml = document_xml.replace(old, new)
            fixes.append("SDD: corrected the Presentation Layer textbox description in Section 2.3.")
        if heading_old in document_xml:
            document_xml = document_xml.replace(heading_old, heading_new, 1)
            fixes.append("SDD: labeled the class diagram section as conceptual.")
        if not fixes:
            return []
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            temp_path = Path(tmp.name)
        with ZipFile(temp_path, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = document_xml.encode("utf-8")
                zout.writestr(item, data)
    shutil.move(str(temp_path), SDD_PATH)
    return fixes


def doc_text(path: Path) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.append("\t".join(cell.text for cell in row.cells))
    with ZipFile(path, "r") as zf:
        chunks.append(zf.read("word/document.xml").decode("utf-8"))
    return "\n".join(chunks)


def validate() -> dict[str, object]:
    frs_text = doc_text(FRS_PATH)
    rtm_doc = Document(RTM_PATH)
    sdd_text = doc_text(SDD_PATH)

    fr_top = sorted(set(re.findall(r"\bFR-\d{2}\b", frs_text)))
    tc_ids = [row.cells[6].text.strip() for row in rtm_doc.tables[0].rows[1:]]
    services = set()
    for row in Document(SDD_PATH).tables[1].rows[1:]:
        services.add(row.cells[0].text.strip())
    rtm_services = set()
    for row in rtm_doc.tables[0].rows[1:]:
        for name in row.cells[3].text.split(";"):
            if name.strip():
                rtm_services.add(name.strip())

    uml_files = {p.name for p in (DOCS / "UML").rglob("*.puml")}
    referenced_uml = []
    for row in rtm_doc.tables[0].rows[1:]:
        referenced_uml.extend([v.strip() for v in row.cells[5].text.split(";") if v.strip()])

    sql_tables = {
        "ApplicationTypes",
        "Applications",
        "Countries",
        "DetainedLicenses",
        "Drivers",
        "InternationalLicenses",
        "LicenseClasses",
        "Licenses",
        "LocalDrivingLicenseApplications",
        "People",
        "TestAppointments",
        "TestTypes",
        "Tests",
        "Users",
    }
    referenced_tables = set()
    for row in rtm_doc.tables[0].rows[1:]:
        referenced_tables.update(v.strip() for v in row.cells[4].text.split(";") if v.strip())

    return {
        "fr_01_to_12_present": fr_top == [f"FR-{i:02d}" for i in range(1, 13)],
        "tc_ids": tc_ids,
        "tc_01_to_20_present": tc_ids == [f"TC-{i:02d}" for i in range(1, 21)],
        "rtm_services_not_in_sdd": sorted(rtm_services - services),
        "missing_uml_files": sorted(set(referenced_uml) - uml_files),
        "invalid_database_tables": sorted(referenced_tables - sql_tables),
        "bad_presentation_text_present": "The Business Layer contains services responsible for business rules and operational workflows, including application types, license classes, test types, countries, and audit logging services." in sdd_text,
        "correct_presentation_text_present": "The Presentation Layer is responsible for rendering the graphical user interface and handling user interactions. It forwards user requests to the Business Layer and displays processed results." in sdd_text,
        "session_protection_present": "Session Protection" in sdd_text,
        "conceptual_class_diagram_present": "3.3 Conceptual Class Diagram" in sdd_text,
    }


def write_report(fixes: list[str], validation: dict[str, object]) -> None:
    remaining = []
    if validation["rtm_services_not_in_sdd"]:
        remaining.append(f"RTM references services not listed in SDD: {validation['rtm_services_not_in_sdd']}")
    if validation["missing_uml_files"]:
        remaining.append(f"RTM references UML files not found by filename: {validation['missing_uml_files']}")
    if validation["invalid_database_tables"]:
        remaining.append(f"RTM references database tables not found in script.sql table list: {validation['invalid_database_tables']}")
    if validation["bad_presentation_text_present"]:
        remaining.append("Incorrect Presentation Layer wording remains in the SDD.")
    if not remaining:
        remaining.append("No remaining clear issues identified in the reviewed scope.")

    lines = [
        "# Final Documentation Quality Check Report",
        "",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Files Reviewed",
        "",
        "- `docs/Functional_Requirements_Specification_FRS.docx`",
        "- `docs/Requirements_Traceability_Matrix.docx`",
        "- `docs/System_Design_Document_SDD.docx`",
        "- Cross-check references: `docs/User_Stories.docx`, `docs/QA_Test_Plan.docx`, `script.sql`, and UML `.puml` files under `docs/UML`.",
        "",
        "## Issues Found",
        "",
        "- Minor FRS wording and punctuation inconsistencies in several requirement rows.",
        "- RTM TC-01 through TC-20 rows were present, but several mappings did not match the QA test plan scenarios.",
        "- RTM rows used service names not listed in the SDD service layer.",
        "- SDD Section 2.3 Presentation Layer textbox incorrectly described the Business Layer.",
        "- SDD security table contained duplicate `Session Management` row labels.",
        "- SDD class diagram heading needed a clearer conceptual label.",
        "",
        "## Issues Fixed",
        "",
    ]
    lines.extend(f"- {fix}" for fix in fixes)
    lines.extend(
        [
            "",
            "## Validation Summary",
            "",
            f"- FRS FR-01 through FR-12 present: `{validation['fr_01_to_12_present']}`",
            f"- RTM TC-01 through TC-20 present and ordered: `{validation['tc_01_to_20_present']}`",
            f"- RTM services not listed in SDD service layer: `{validation['rtm_services_not_in_sdd']}`",
            f"- RTM UML filenames missing from `docs/UML`: `{validation['missing_uml_files']}`",
            f"- RTM database table names not in SQL table list: `{validation['invalid_database_tables']}`",
            f"- Correct SDD Presentation Layer wording present: `{validation['correct_presentation_text_present']}`",
            f"- Duplicate session row resolved with `Session Protection`: `{validation['session_protection_present']}`",
            f"- Class diagram labeled conceptual: `{validation['conceptual_class_diagram_present']}`",
            "",
            "## Remaining Issues",
            "",
        ]
    )
    lines.extend(f"- {item}" for item in remaining)
    lines.extend(
        [
            "",
            "## Scope Confirmation",
            "",
            "- No database schema changes were made.",
            "- No architecture redesign was performed.",
            "- Existing document structure was preserved; changes were limited to consistency, traceability, wording, and obvious documentation issues.",
        ]
    )
    REPORT_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    fixes: list[str] = []
    fixes.extend(update_frs())
    fixes.extend(update_rtm())
    fixes.extend(update_sdd_tables())
    fixes.extend(patch_sdd_textbox())
    validation = validate()
    write_report(fixes, validation)
    for key, value in validation.items():
        print(f"{key}={value}")


if __name__ == "__main__":
    main()
