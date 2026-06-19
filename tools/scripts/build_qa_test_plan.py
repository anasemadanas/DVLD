from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "docs" / "QA_Test_Plan.docx"
REPORT = ROOT / "DOCUMENTATION_IMPROVEMENT_REPORT.md"
TODAY = "2026-06-16"


def set_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(6)

    tokens = {
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, widths=None, font_size=9) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                if index < len(row.cells):
                    row.cells[index].width = Inches(width)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)

    header = table.rows[0]
    for cell in header.cells:
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("0B2545")


def add_table(doc: Document, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table, widths=widths, font_size=font_size)
    return table


def add_bullets(doc: Document, items) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_document() -> None:
    doc = Document()
    set_styles(doc)
    configure_page(doc)

    props = doc.core_properties
    props.title = "DVLD System - QA Test Plan"
    props.subject = "Quality assurance test planning for the DVLD System"
    props.author = "DVLD Documentation Team"
    props.keywords = "DVLD, QA, test plan, validation, testing"

    doc.add_paragraph("DVLD System - QA Test Plan", style="Title")
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(
        "Company-grade quality assurance plan defining how the DVLD System will be tested and validated before delivery."
    )
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string("44546A")
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Document date: {TODAY}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor.from_string("667085")

    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph(
        "This QA Test Plan defines the strategy used to test and validate the DVLD System before delivery, demonstration, or formal handoff. "
        "It establishes the testing scope, test objectives, test types, environment, roles, entry criteria, exit criteria, and sample test cases required to confirm business readiness."
    )

    doc.add_paragraph("Scope", style="Heading 1")
    add_bullets(
        doc,
        [
            "User Management",
            "Person Management",
            "License Applications",
            "License Eligibility Validation",
            "Test Scheduling",
            "Test Result Recording",
            "License Issuance",
            "License Renewal",
            "Lost and Damaged License Replacement",
            "Detention and Release",
            "International License Issuance",
            "Reports and Search",
            "Audit Logging",
        ],
    )

    doc.add_paragraph("Test Objectives", style="Heading 1")
    add_bullets(
        doc,
        [
            "Verify that functional requirements are implemented correctly.",
            "Verify that DVLD business rules are enforced consistently.",
            "Verify that validation rules prevent invalid operations and incomplete processing.",
            "Verify that user permissions restrict unauthorized access to sensitive features.",
            "Verify that database records remain consistent after successful and rejected workflows.",
            "Verify that day-to-day user workflows operate correctly for employees, administrators, and test officers.",
        ],
    )

    doc.add_paragraph("Test Types", style="Heading 1")
    add_table(
        doc,
        ["Test Type", "Purpose", "Typical Focus"],
        [
            (
                "Functional Testing",
                "Confirms that each DVLD feature behaves according to the approved requirements and workflow rules.",
                "Core services, business workflows, confirmations, and expected outputs.",
            ),
            (
                "Validation Testing",
                "Confirms that business validations and input rules reject invalid, incomplete, or duplicate operations.",
                "Duplicate prevention, age rules, license status rules, required-field rules, and test sequencing.",
            ),
            (
                "Security Testing",
                "Confirms that authentication and authorization protect restricted operations and administrative functions.",
                "Login, permissions, inactive users, role-based restrictions, and controlled access.",
            ),
            (
                "Usability Testing",
                "Confirms that employees can complete key workflows clearly, efficiently, and with understandable feedback.",
                "Navigation, forms, messages, search behavior, and screen readability.",
            ),
            (
                "Integration Testing",
                "Confirms that UI, service logic, repositories, and SQL Server persistence work together without data loss or workflow breaks.",
                "Cross-layer interactions and multi-step business services.",
            ),
            (
                "Regression Testing",
                "Confirms that previously working features remain stable after fixes, enhancements, or documentation-aligned changes.",
                "High-risk workflows and existing validated scenarios.",
            ),
            (
                "Database Testing",
                "Confirms that inserts, updates, lookups, constraints, and relationship behavior remain correct and consistent.",
                "PK/FK integrity, unique rules, stored records, and transaction consistency.",
            ),
            (
                "Error Handling Testing",
                "Confirms that the system responds safely and clearly to invalid input, failed operations, and blocked workflows.",
                "Validation messages, rejected actions, state protection, and user guidance.",
            ),
        ],
        widths=[1.45, 3.6, 4.35],
    )

    doc.add_paragraph("Test Environment", style="Heading 1")
    add_table(
        doc,
        ["Environment Item", "Configuration"],
        [
            ("Operating System", "Windows"),
            ("Application Type", "C# desktop application"),
            ("Database", "Microsoft SQL Server"),
            (
                "Test Data",
                "Sample applicants, users, licenses, applications, tests, license classes, and detention records representing valid and invalid scenarios.",
            ),
            (
                "Required Tools",
                "Visual Studio, SQL Server Management Studio, and PlantUML or related documentation tools where supporting document review is required.",
            ),
        ],
        widths=[1.9, 7.8],
    )

    doc.add_paragraph("Test Roles and Responsibilities", style="Heading 1")
    add_table(
        doc,
        ["Role", "Responsibilities"],
        [
            (
                "QA Engineer",
                "Prepare and maintain test cases, execute tests, record results, log defects, verify fixes, and contribute to the final test summary.",
            ),
            (
                "Developer",
                "Clarify implementation behavior, support defect analysis, fix reported issues, and confirm technical resolution readiness.",
            ),
            (
                "System Administrator",
                "Prepare user accounts, maintain the test environment, manage application access, and support database and configuration readiness.",
            ),
            (
                "Licensing Employee / Business User",
                "Validate that workflows reflect DVLD operational practice and confirm business correctness during user acceptance-style review.",
            ),
            (
                "Test Officer",
                "Validate test scheduling and result-recording workflows, confirm operational accuracy, and review pass/fail processing behavior.",
            ),
        ],
        widths=[2.2, 7.5],
    )

    doc.add_paragraph("Entry Criteria", style="Heading 1")
    add_bullets(
        doc,
        [
            "Requirements documents are available.",
            "Functional requirements are approved.",
            "The test environment is prepared.",
            "The database is available.",
            "Required test data is prepared.",
            "The application build is ready for testing.",
        ],
    )

    doc.add_paragraph("Exit Criteria", style="Heading 1")
    add_bullets(
        doc,
        [
            "Critical defects are resolved.",
            "Major business workflows pass testing.",
            "Required test cases are executed.",
            "Regression testing is completed.",
            "Test results are documented.",
            "The system is approved for delivery or demonstration.",
        ],
    )

    doc.add_paragraph("Test Case Format", style="Heading 1")
    add_table(
        doc,
        ["Field", "Description"],
        [
            ("Test Case ID", "Unique identifier assigned to the test case."),
            ("Related Requirement ID", "Functional requirement or grouped requirement covered by the test case."),
            ("Test Scenario", "Short business-oriented statement of what is being tested."),
            ("Preconditions", "Required system state, user role, or data setup before execution."),
            ("Test Steps", "Ordered actions the tester performs to execute the scenario."),
            ("Expected Result", "The correct system outcome if the feature works as intended."),
            ("Actual Result", "Execution outcome recorded by QA during test execution."),
            ("Status", "Execution state such as Not Executed, Passed, Failed, or Blocked."),
            ("Notes", "Supporting observations, evidence references, or clarification notes."),
        ],
        widths=[2.0, 7.7],
    )

    doc.add_paragraph("Sample Test Cases", style="Heading 1")
    sample_cases = [
        (
            "TC-01",
            "FR-01",
            "Login with valid credentials",
            "Active user account exists and user has assigned permissions.",
            "Open login screen; enter valid username and password; select Sign In.",
            "User is authenticated successfully and the dashboard opens with permitted modules.",
            "To be recorded during execution.",
            "Not Executed",
            "Evidence reference to be added during testing.",
        ),
        (
            "TC-02",
            "FR-01",
            "Login with invalid credentials",
            "Login screen is available.",
            "Open login screen; enter invalid username or password; submit the request.",
            "Access is denied and a clear authentication error message is displayed without creating a session.",
            "To be recorded during execution.",
            "Not Executed",
            "Negative authentication scenario.",
        ),
        (
            "TC-03",
            "FR-02",
            "Add new person",
            "Authorized user is signed in and country reference data exists.",
            "Open Person Management; enter a new person's details; save the record.",
            "Person record is saved successfully with required details and becomes searchable by National Number.",
            "To be recorded during execution.",
            "Not Executed",
            "Positive master-data scenario.",
        ),
        (
            "TC-04",
            "FR-02",
            "Reject duplicate national number",
            "A person record already exists with the target National Number.",
            "Open Person Management; enter a new person using the existing National Number; save the record.",
            "The system rejects the record and displays a validation message indicating that the National Number must be unique.",
            "To be recorded during execution.",
            "Not Executed",
            "Uniqueness validation scenario.",
        ),
        (
            "TC-05",
            "FR-04; FR-05",
            "Create license application",
            "Eligible applicant exists and the user has application-management permission.",
            "Open Applications; select the applicant, application type, and license class; save the application.",
            "A new application and linked local driving license application are created successfully.",
            "To be recorded during execution.",
            "Not Executed",
            "Positive application-creation scenario.",
        ),
        (
            "TC-06",
            "FR-04; FR-05",
            "Reject duplicate active application",
            "Applicant already has an active application for the same service and class.",
            "Attempt to create another active application of the same type for the same applicant.",
            "The system blocks the operation and displays a duplicate active application validation message.",
            "To be recorded during execution.",
            "Not Executed",
            "Duplicate-prevention scenario.",
        ),
        (
            "TC-07",
            "FR-07",
            "Schedule vision test",
            "Application exists, required fees are recorded, and no active duplicate vision appointment exists.",
            "Open Test Scheduling; choose Vision Test; select date and details; save the appointment.",
            "Vision test appointment is created successfully and linked to the correct application.",
            "To be recorded during execution.",
            "Not Executed",
            "First test-stage scheduling scenario.",
        ),
        (
            "TC-08",
            "FR-07",
            "Record test result",
            "Scheduled test appointment exists and authorized test officer is signed in.",
            "Open the appointment; record pass or fail result and score where applicable; save the result.",
            "Result is stored successfully, the appointment is completed, and score validation rules are enforced.",
            "To be recorded during execution.",
            "Not Executed",
            "Covers written-score or pass/fail recording.",
        ),
        (
            "TC-09",
            "FR-07; FR-09.1",
            "Prevent license issuance before all tests pass",
            "Application exists but one or more required tests are missing or failed.",
            "Open License Issuance; enter the application number; attempt to issue the license.",
            "The system blocks issuance and displays a message stating that all required tests must be passed first.",
            "To be recorded during execution.",
            "Not Executed",
            "Critical pre-issuance validation scenario.",
        ),
        (
            "TC-10",
            "FR-03; FR-05; FR-09.1",
            "Issue driving license after all conditions pass",
            "Applicant is eligible, all required tests are passed, and application is ready for issuance.",
            "Open License Issuance; enter the application number; validate eligibility; confirm issuance.",
            "Driver record is created if needed, a unique license number is generated, the license is saved, and confirmation is displayed.",
            "To be recorded during execution.",
            "Not Executed",
            "Core business workflow.",
        ),
        (
            "TC-11",
            "FR-09.2",
            "Renew expired license",
            "Expired license exists and renewal prerequisites are satisfied.",
            "Open License Renewal; search for the license; confirm required checks; process renewal.",
            "Renewal is completed successfully and the old license is replaced or marked inactive according to business rules.",
            "To be recorded during execution.",
            "Not Executed",
            "Renewal validation scenario.",
        ),
        (
            "TC-12",
            "FR-09.3",
            "Replace lost license",
            "Valid license exists and is not detained.",
            "Open Lost License Replacement; locate the license; confirm details; issue replacement.",
            "Replacement license is issued successfully and replacement records are stored.",
            "To be recorded during execution.",
            "Not Executed",
            "Positive replacement scenario.",
        ),
        (
            "TC-13",
            "FR-09.4",
            "Replace damaged license",
            "Damaged license exists and surrender condition can be satisfied.",
            "Open Damaged License Replacement; locate the license; confirm surrender and issue replacement.",
            "Replacement license is issued successfully and damaged-license replacement details are recorded.",
            "To be recorded during execution.",
            "Not Executed",
            "Positive damaged-replacement scenario.",
        ),
        (
            "TC-14",
            "FR-10",
            "Detain license",
            "Active license exists and the user is authorized to perform detention.",
            "Open Detention; enter license details, fine amount, and reason; save the detention record.",
            "License is detained successfully and detention date, fine, reason, and responsible user are recorded.",
            "To be recorded during execution.",
            "Not Executed",
            "Enforces detention record completeness.",
        ),
        (
            "TC-15",
            "FR-09.5; FR-10",
            "Release detained license after fine payment",
            "License has an active detention and fine payment has been completed.",
            "Open License Release; verify fine payment; create release application; confirm release.",
            "Detention is released successfully and the release date and linked application are recorded.",
            "To be recorded during execution.",
            "Not Executed",
            "Post-payment release scenario.",
        ),
        (
            "TC-16",
            "FR-09.6",
            "Issue international license for valid Class 3 license",
            "Driver holds a valid active non-detained Class 3 local license.",
            "Open International License Issuance; search for the driver; confirm eligibility; issue the license.",
            "International license is issued successfully and linked to the correct driver and local license.",
            "To be recorded during execution.",
            "Not Executed",
            "Positive international-license scenario.",
        ),
        (
            "TC-17",
            "FR-09.6",
            "Reject international license for detained or expired license",
            "Driver's local license is detained, expired, or otherwise ineligible.",
            "Attempt to issue an international license for the ineligible local license.",
            "The system rejects issuance and displays a clear eligibility or status validation message.",
            "To be recorded during execution.",
            "Not Executed",
            "Negative international-license scenario.",
        ),
        (
            "TC-18",
            "FR-02; FR-09.7",
            "Search records by national number",
            "Authorized user is signed in and target person record exists.",
            "Open Search or Inquiry; enter the National Number; run the search.",
            "Matching person, application, driver, or license records are returned accurately.",
            "To be recorded during execution.",
            "Not Executed",
            "Inquiry accuracy scenario.",
        ),
        (
            "TC-19",
            "FR-12",
            "Filter reports by status, date, and license class",
            "Search and reporting data exists for multiple statuses, dates, and license classes.",
            "Open reporting or inquiry screen; apply status, date, and license class filters; run the report.",
            "Only records matching the selected filters are displayed in the result set.",
            "To be recorded during execution.",
            "Not Executed",
            "Reporting filter scenario.",
        ),
        (
            "TC-20",
            "FR-11",
            "Verify audit log records operation user and date",
            "Auditable business operation can be completed by an authorized user.",
            "Perform a business transaction; open audit logs or inspect stored audit data for the transaction.",
            "Audit information includes the operation type, responsible user, and operation date for the completed action.",
            "To be recorded during execution.",
            "Not Executed",
            "Audit-traceability scenario.",
        ),
    ]
    add_table(
        doc,
        [
            "Test Case ID",
            "Related Requirement ID",
            "Test Scenario",
            "Preconditions",
            "Test Steps",
            "Expected Result",
            "Actual Result",
            "Status",
            "Notes",
        ],
        sample_cases,
        widths=[0.6, 0.9, 1.15, 1.15, 1.75, 1.5, 1.1, 0.7, 0.75],
        font_size=8,
    )

    doc.add_paragraph("Defect Severity Levels", style="Heading 1")
    add_table(
        doc,
        ["Severity", "Description"],
        [
            (
                "Critical",
                "Stops delivery or demonstration because a core workflow fails, data integrity is compromised, or a major security or system-stability issue exists.",
            ),
            (
                "High",
                "A major business function is broken or unreliable, but limited workaround options may exist.",
            ),
            (
                "Medium",
                "A non-blocking functional or usability issue affects efficiency, consistency, or a secondary workflow.",
            ),
            (
                "Low",
                "Minor issue with limited business impact, such as cosmetic wording, formatting, or low-risk convenience behavior.",
            ),
        ],
        widths=[1.6, 8.1],
    )

    doc.add_paragraph("Test Deliverables", style="Heading 1")
    add_table(
        doc,
        ["Deliverable", "Purpose"],
        [
            ("QA Test Plan", "Defines QA strategy, scope, responsibilities, and validation readiness."),
            ("Test Cases", "Provide executable scenarios linked to requirements and business rules."),
            ("Test Execution Results", "Capture actual outcomes, status, evidence, and execution notes during testing."),
            ("Defect Reports", "Record identified issues, severity, impact, reproduction details, and fix status."),
            ("Final Test Summary", "Summarizes coverage, outcomes, residual risks, and readiness for delivery or demonstration."),
        ],
        widths=[2.0, 7.7],
    )

    doc.add_paragraph("Traceability", style="Heading 1")
    doc.add_paragraph(
        "Each executed test case should be linked to the related functional requirements, user stories, use cases, and business rules so QA coverage can be demonstrated clearly. "
        "The QA Test Plan works alongside the DVLD requirements and traceability documents to support implementation review, defect analysis, and delivery readiness."
    )
    add_bullets(
        doc,
        [
            "Functional Requirements",
            "User Stories",
            "Use Cases",
            "Business Rules",
        ],
    )

    doc.save(OUTPUT)


def update_report() -> None:
    text = REPORT.read_text(encoding="utf-8")
    old_summary = """## QA Test Plan Summary\n\n`QA_Test_Plan.docx` was created with the following sections:\n\n- Introduction\n- Scope\n- Test Objectives\n- Test Types\n- Functional Testing\n- Validation Testing\n- Security Testing\n- Usability Testing\n- Regression Testing\n- Test Environment\n- Test Roles\n- Entry and Exit Criteria\n- Sample Test Cases\n\nSample test cases were included for:\n\n- Login\n- Create Application\n- Schedule Test\n- Record Test Result\n- Issue License\n- Renew License\n- Replace License\n- Detain License\n- Release License\n- Search Records\n"""
    new_summary = """## QA Test Plan Summary\n\n`docs/QA_Test_Plan.docx` was recreated as `DVLD System - QA Test Plan` and now includes the following sections:\n\n- Introduction\n- Scope\n- Test Objectives\n- Test Types\n- Test Environment\n- Test Roles and Responsibilities\n- Entry Criteria\n- Exit Criteria\n- Test Case Format\n- Sample Test Cases\n- Defect Severity Levels\n- Test Deliverables\n- Traceability\n\nThe plan now includes 20 professional sample test cases covering authentication, person management, license applications, validation failures, test scheduling, test result recording, license issuance, renewal, replacement, detention, release, international licensing, search, reporting, and audit logging.\n"""
    if old_summary in text:
        text = text.replace(old_summary, new_summary)
    elif "## QA Test Plan Summary" not in text:
        insert_at = text.find("## UML Catalog Summary")
        if insert_at != -1:
            text = text[:insert_at] + new_summary + "\n" + text[insert_at:]
        else:
            text += "\n\n" + new_summary
    REPORT.write_text(text, encoding="utf-8")


if __name__ == "__main__":
    build_document()
    update_report()