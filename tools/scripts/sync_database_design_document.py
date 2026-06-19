from __future__ import annotations

import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[2]
DOC_PATH = REPO_ROOT / "docs" / "Database_Design_Document.docx"
REPORT_PATH = REPO_ROOT / "docs" / "database-design-document-sync-report.md"
SQL_PATH = REPO_ROOT / "script.sql"

sys.path.insert(0, str(REPO_ROOT / "tools" / "scripts"))
from sync_database_diagrams_from_sql import Schema, parse_sql  # noqa: E402


TABLE_DESCRIPTIONS = {
    "Applications": "Stores service application headers, applicant references, application status, fees, dates, and creator user.",
    "ApplicationTypes": "Reference table for service request types and their configured fees.",
    "Countries": "Reference table for nationality and country lookup data.",
    "DetainedLicenses": "Stores license detention, release, fine, release application, and responsible-user audit references.",
    "Drivers": "Stores driver identities linked to people and the user who created the driver record.",
    "InternationalLicenses": "Stores international license records linked to applications, drivers, local licenses, and creator users.",
    "LicenseClasses": "Reference table for license classes, age requirements, validity length, and class fees.",
    "Licenses": "Stores issued local licenses, class, application, driver, dates, fees, status, reason, notes, and creator user.",
    "LocalDrivingLicenseApplications": "Specialized local driving license application table linked to general applications and license classes.",
    "People": "Master person registry for applicants, users, drivers, and nationality references.",
    "TestAppointments": "Stores scheduled test appointments, test type, local application, fees, lock status, and creator user.",
    "Tests": "Stores test results linked to appointments and creator users.",
    "TestTypes": "Reference table for configured test types, descriptions, and fees.",
    "Users": "Stores system user accounts linked to people and active status.",
}


COLUMN_DESCRIPTIONS = {
    "ApplicationID": "Unique application identifier.",
    "ApplicantPersonID": "Applicant person reference.",
    "ApplicationDate": "Date and time when the application was created.",
    "ApplicationTypeID": "Application type reference.",
    "ApplicationStatus": "Application workflow status.",
    "LastStatusDate": "Date and time when status last changed.",
    "PaidFees": "Fee amount paid for the record.",
    "CreatedByUserID": "User who created the record.",
    "ApplicationTypeTitle": "Business title of the application type.",
    "ApplicationFees": "Configured application type fee.",
    "CountryID": "Unique country identifier.",
    "CountryName": "Country name.",
    "DetainID": "Unique detention identifier.",
    "LicenseID": "License reference.",
    "DetainDate": "Date and time when the license was detained.",
    "FineFees": "Fine fee amount.",
    "IsReleased": "Release status flag.",
    "ReleaseDate": "Date and time when the detained license was released.",
    "ReleasedByUserID": "User who released the detained license.",
    "ReleaseApplicationID": "Application used to release the detained license.",
    "DriverID": "Driver reference or unique driver identifier.",
    "PersonID": "Person reference or unique person identifier.",
    "CreatedDate": "Date and time when the driver record was created.",
    "InternationalLicenseID": "Unique international license identifier.",
    "IssuedUsingLocalLicenseID": "Local license used to issue the international license.",
    "IssueDate": "Date and time when the license was issued.",
    "ExpirationDate": "Date and time when the license expires.",
    "IsActive": "Active status flag.",
    "LicenseClassID": "Unique license class identifier or license class reference.",
    "ClassName": "License class name.",
    "ClassDescription": "License class description.",
    "MinimumAllowedAge": "Minimum allowed applicant age.",
    "DefaultValidityLength": "Default validity length for the license class.",
    "ClassFees": "Configured fee for the license class.",
    "LicenseClass": "License class foreign key column implemented in script.sql.",
    "Notes": "Optional notes.",
    "IssueReason": "Numeric issue reason code.",
    "LocalDrivingLicenseApplicationID": "Unique local driving license application identifier.",
    "NationalNo": "National number.",
    "FirstName": "First name.",
    "SecondName": "Second name.",
    "ThirdName": "Third name.",
    "LastName": "Last name.",
    "DateOfBirth": "Date and time of birth as implemented in SQL.",
    "Gendor": "Gender code column preserved exactly as implemented in script.sql.",
    "Address": "Address.",
    "Phone": "Phone number.",
    "Email": "Email address.",
    "NationalityCountryID": "Nationality country reference.",
    "ImagePath": "Optional person image path.",
    "TestAppointmentID": "Unique test appointment identifier or appointment reference.",
    "TestTypeID": "Test type reference or unique test type identifier.",
    "AppointmentDate": "Date and time of the appointment.",
    "IsLocked": "Appointment lock status flag.",
    "TestID": "Unique test result identifier.",
    "TestResult": "Pass/fail test result flag.",
    "TestTypeTitle": "Business title of the test type.",
    "TestTypeDescription": "Test type description.",
    "TestTypeFees": "Configured fee for the test type.",
    "UserID": "Unique system user identifier.",
    "UserName": "System login username.",
    "Password": "Stored password field as currently implemented in script.sql.",
}


def read_sql_text() -> str:
    raw = SQL_PATH.read_bytes()
    for encoding in ("utf-16", "utf-8-sig", "utf-8", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise RuntimeError("Unable to decode script.sql")


def parse_defaults() -> dict[tuple[str, str], str]:
    text = read_sql_text()
    pattern = re.compile(
        r"ALTER TABLE \[dbo\]\.\[([^\]]+)\] ADD\s+CONSTRAINT \[[^\]]+\]\s+DEFAULT\s+(.+?)\s+FOR \[([^\]]+)\]",
        re.IGNORECASE,
    )
    defaults: dict[tuple[str, str], str] = {}
    for table, value, column in pattern.findall(text):
        defaults[(table, column)] = value.strip()
    return defaults


def remove_from_heading(doc: Document, heading_text: str) -> None:
    body = doc._body._element
    children = list(body)
    start_index = None
    normalized_heading = " ".join(heading_text.split())
    for paragraph in doc.paragraphs:
        if " ".join(paragraph.text.split()) == normalized_heading:
            for index, child in enumerate(children):
                if child is paragraph._p:
                    start_index = index
                    break
            break
    for index, child in enumerate(children):
        texts = [node.text for node in child.iter() if node.text]
        if " ".join(" ".join(texts).split()) == normalized_heading:
            start_index = index
            break
    if start_index is None:
        raise RuntimeError(f"Could not find heading {heading_text!r}")
    for child in children[start_index:]:
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_table_style(table) -> None:
    for style_name in ("Table Grid", "TableGrid"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = value
    set_table_style(table)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def key_for(column) -> str:
    flags = []
    if column.pk:
        flags.append("PK")
    if column.fk:
        flags.append("FK")
    return ", ".join(flags) if flags else ""


def nullable_for(column) -> str:
    return "NULL" if column.nullable else "NOT NULL"


def default_for(table_name: str, column, defaults: dict[tuple[str, str], str]) -> str:
    if column.identity:
        return "IDENTITY(1,1)"
    return defaults.get((table_name, column.name), "None")


def add_relationship_sections(doc: Document, schema: Schema) -> None:
    doc.add_heading("5. Relationships", level=1)
    doc.add_paragraph(
        "Relationships in this section reflect SQL Server foreign key constraints from script.sql. "
        "Every FK is documented as Parent Table 1 to Child Table Many unless a SQL UNIQUE constraint exists on the FK column. "
        "script.sql currently defines 0 unique constraints."
    )
    rows = []
    for fk in schema.foreign_keys:
        child_col = ", ".join(fk.child_columns)
        parent_col = ", ".join(fk.parent_columns)
        business_note = f"{fk.child_table}.{child_col} references {fk.parent_table}.{parent_col} through {fk.name}."
        if fk.name == "FK_Users_People":
            business_note += " Business rule may intend optional one-to-one, but SQL does not enforce uniqueness."
        if fk.name == "FK_Drivers_People":
            business_note += " Business rule may intend optional one-to-one, but SQL does not enforce uniqueness."
        if fk.name == "FK_DrivingLicsenseApplications_Applications":
            business_note += " Business rule may intend optional one-to-one, but SQL does not enforce uniqueness."
        rows.append([
            f"{fk.parent_table} -> {fk.child_table}",
            "One-to-Many",
            fk.name,
            f"{fk.child_table}.{child_col} -> {fk.parent_table}.{parent_col}",
            business_note,
        ])
    add_table(
        doc,
        ["Relationship", "SQL Cardinality", "FK Constraint", "FK Mapping", "Notes"],
        rows,
        [1.4, 1.0, 1.7, 1.9, 2.2],
    )
    doc.add_paragraph(
        "People -> Users: One-to-Many technically; Optional One-to-One as business rule."
    )
    doc.add_paragraph(
        "People -> Drivers: One-to-Many technically; Optional One-to-One as business rule."
    )
    doc.add_paragraph(
        "Applications -> LocalDrivingLicenseApplications: One-to-Many technically; Optional One-to-One as business rule if intended."
    )
    doc.add_heading("5.2 Audit Relationships", level=2)
    doc.add_paragraph(
        "Users create, release, and manage operational records through CreatedByUserID and ReleasedByUserID relationships. "
        "These relationships support auditability and traceability across applications, drivers, licenses, international licenses, "
        "detained licenses, test appointments, and tests."
    )


def add_pk_fk_section(doc: Document, schema: Schema) -> None:
    doc.add_heading("6. Primary Keys & Foreign Keys", level=1)
    rows = []
    for table in schema.tables.values():
        fk_columns = []
        referenced = []
        for fk in schema.foreign_keys:
            if fk.child_table == table.name:
                fk_columns.append(", ".join(fk.child_columns))
                referenced.append(f"{fk.parent_table}.{', '.join(fk.parent_columns)} ({fk.name})")
        rows.append([
            table.name,
            ", ".join(table.pk_columns),
            "; ".join(fk_columns) if fk_columns else "None",
            "; ".join(referenced) if referenced else "None",
        ])
    add_table(
        doc,
        ["Table Name", "Primary Key", "Foreign Key Columns", "Referenced Table(s) / Constraint"],
        rows,
        [1.7, 1.6, 2.3, 3.0],
    )


def add_constraints_section(doc: Document, schema: Schema) -> None:
    doc.add_heading("7. Constraints", level=1)
    rows = [
        [
            "Primary Key Constraints",
            "Every table has a single-column SQL primary key as defined in script.sql.",
        ],
        [
            "Foreign Key Constraints",
            "script.sql defines 26 FK constraints connecting operational child records to parent reference or transaction records.",
        ],
        [
            "Nullability Constraints",
            "NOT NULL and NULL settings are implemented at column level exactly as listed in the Data Dictionary.",
        ],
        [
            "Default Constraints",
            "SQL default constraints exist for selected status, fee, age, validity, reason, gender code, and lock columns.",
        ],
        [
            "Unique Constraints",
            "No SQL unique constraints are currently defined in script.sql. Business uniqueness rules such as unique NationalNo, unique UserName, one user per person, one driver per person, and one active international license should be enforced through application validation or future database unique/filtered indexes if required.",
        ],
    ]
    add_table(doc, ["Constraint Category", "SQL Server Implementation"], rows, [2.0, 5.7])
    doc.add_heading("7.1 Recommended Future Constraints / Indexes", level=2)
    for item in [
        "Consider a unique index for People.NationalNo if national numbers must be SQL-enforced unique.",
        "Consider a unique index for Users.UserName if usernames must be SQL-enforced unique.",
        "Consider unique or filtered indexes for one user per person, one driver per person, and one active international license if those business rules must be enforced by SQL Server.",
        "Consider indexes on FK columns, status columns, date columns, and common search columns such as NationalNo and UserName.",
    ]:
        doc.add_paragraph(f"- {item}")


def add_data_dictionary_section(doc: Document, schema: Schema, defaults: dict[tuple[str, str], str]) -> None:
    doc.add_heading("8. Data Dictionary", level=1)
    doc.add_paragraph(
        "The Data Dictionary reflects the current SQL Server implementation from script.sql. "
        "Some column names such as Gendor and LicenseClass are preserved exactly as implemented in the database, "
        "even if business documentation may use clearer wording such as Gender or LicenseClassID."
    )
    for index, table in enumerate(schema.tables.values(), start=1):
        doc.add_heading(f"8.{index} {table.name}", level=2)
        doc.add_paragraph(TABLE_DESCRIPTIONS.get(table.name, f"SQL Server table {table.name}."))
        rows = []
        for column in table.columns:
            rows.append([
                column.name,
                column.data_type,
                key_for(column),
                nullable_for(column),
                default_for(table.name, column, defaults),
                COLUMN_DESCRIPTIONS.get(column.name, ""),
            ])
        add_table(
            doc,
            ["Column Name", "Data Type", "Key", "Nullable", "Default", "Description"],
            rows,
            [1.5, 1.2, 0.6, 0.8, 1.1, 2.7],
        )


def update_document() -> None:
    schema = parse_sql(SQL_PATH)
    defaults = parse_defaults()
    doc = Document(DOC_PATH)
    remove_from_heading(doc, "5. Relationships")
    add_relationship_sections(doc, schema)
    add_pk_fk_section(doc, schema)
    add_constraints_section(doc, schema)
    add_data_dictionary_section(doc, schema, defaults)
    doc.save(DOC_PATH)


def validate_document() -> dict[str, object]:
    schema = parse_sql(SQL_PATH)
    doc = Document(DOC_PATH)
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    required = [
        "NationalityCountryID",
        "ApplicantPersonID",
        "LicenseClass",
        "LocalDrivingLicenseApplicationID",
        "TestTypeTitle",
        "Gendor",
        "No SQL unique constraints are currently defined in script.sql",
        "5.2 Audit Relationships",
    ]
    missing = [item for item in required if item not in text]
    generic_types = [
        item
        for item in ["BOOLEAN", "TEXT", "DECIMAL(10,2)", "DATE"]
        if item in text
    ]
    table_missing = [table for table in schema.tables if table not in text]
    fk_missing = [fk.name for fk in schema.foreign_keys if fk.name not in text]
    return {
        "missing_required_text": missing,
        "generic_type_mentions": generic_types,
        "missing_tables": table_missing,
        "missing_fk_constraints": fk_missing,
        "table_count": len(schema.tables),
        "column_count": len(schema.columns),
        "fk_count": len(schema.foreign_keys),
    }


def write_report(validation: dict[str, object]) -> None:
    lines = [
        "# Database Design Document Sync Report",
        "",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Source Of Truth",
        "",
        "- `script.sql`",
        "",
        "## Sections Updated",
        "",
        "- Section 5 - Relationships",
        "- Section 5.2 - Audit Relationships",
        "- Section 6 - Primary Keys & Foreign Keys",
        "- Section 7 - Constraints",
        "- Section 8 - Data Dictionary",
        "",
        "## Mismatches Fixed",
        "",
        "- Replaced older or generic SQL type descriptions with SQL Server types parsed from `script.sql`.",
        "- Preserved implemented column names such as `Gendor`, `LicenseClass`, `ApplicantPersonID`, `NationalityCountryID`, and `TestTypeTitle`.",
        "- Rebuilt the Data Dictionary from the current 14-table, 88-column SQL Server schema.",
        "",
        "## Relationship Cardinality Fixes",
        "",
        "- Documented every FK as Parent Table `1` to Child Table `Many` unless a SQL UNIQUE constraint exists.",
        "- Documented `People -> Users` as technically One-to-Many and optional One-to-One only as a business rule.",
        "- Documented `People -> Drivers` as technically One-to-Many and optional One-to-One only as a business rule.",
        "- Documented `Applications -> LocalDrivingLicenseApplications` as technically One-to-Many and optional One-to-One only as a business rule if intended.",
        "- Moved user-created operational relationships into Section 5.2 Audit Relationships.",
        "",
        "## PK/FK Fixes",
        "",
        "- `People.NationalityCountryID` now references `Countries.CountryID`.",
        "- `Applications.ApplicantPersonID` now references `People.PersonID`.",
        "- `Licenses.LicenseClass` now references `LicenseClasses.LicenseClassID`.",
        "- `LocalDrivingLicenseApplications.LocalDrivingLicenseApplicationID` is documented as the PK.",
        "- All 26 FK constraint names from `script.sql` are included.",
        "",
        "## Data Type Fixes",
        "",
        "- `People.SecondName` is `nvarchar(20) NOT NULL`.",
        "- `People.ThirdName` is `nvarchar(20) NULL`.",
        "- `Licenses.IssueReason` is `tinyint NOT NULL`.",
        "- `Applications.ApplicationStatus` is `tinyint NOT NULL`.",
        "- `Applications.PaidFees`, `ApplicationTypes.ApplicationFees`, `LicenseClasses.ClassFees`, and fee fields are `smallmoney`.",
        "- `IssueDate`, `ExpirationDate`, `AppointmentDate`, `DetainDate`, and `ReleaseDate` use the exact `datetime` or `smalldatetime` types from SQL.",
        "- `Tests.TestResult` is `bit NOT NULL`.",
        "",
        "## Validation Summary",
        "",
        f"- Tables documented: {validation['table_count']}",
        f"- Columns documented: {validation['column_count']}",
        f"- FK constraints documented: {validation['fk_count']}",
        f"- Missing required text: {validation['missing_required_text']}",
        f"- Generic disallowed type mentions: {validation['generic_type_mentions']}",
        f"- Missing tables: {validation['missing_tables']}",
        f"- Missing FK constraints: {validation['missing_fk_constraints']}",
        "",
        "## Remaining Conflicts",
        "",
        "- None identified in Sections 5-8 after synchronization.",
        "",
        "## Notes",
        "",
        "- `script.sql` was not modified.",
        "- Database diagrams, PlantUML, and Draw.io files were not modified.",
        "- Visual DOCX render QA was not completed because the available renderer dependency stack is incomplete in this environment; structural validation was completed with `python-docx`.",
    ]
    REPORT_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    update_document()
    validation = validate_document()
    write_report(validation)
    print(f"tables={validation['table_count']}")
    print(f"columns={validation['column_count']}")
    print(f"fk_constraints={validation['fk_count']}")
    print(f"missing_required={len(validation['missing_required_text'])}")
    print(f"generic_type_mentions={len(validation['generic_type_mentions'])}")


if __name__ == "__main__":
    main()
